import pandas as pd
from docx import Document
import os

# Verificar se o arquivo existe
def verificar_arquivo(caminho):
    if not os.path.exists(caminho):
        print(f"Erro: O arquivo '{caminho}' não foi encontrado.")
        return False
    return True

# Carregar as notas fiscais de um arquivo CSV
def carregar_nota_fiscal(nota_fiscal_csv):
    try:
        if verificar_arquivo(nota_fiscal_csv):
            # Usar on_bad_lines='skip' para ignorar linhas problemáticas
            df = pd.read_csv(nota_fiscal_csv, on_bad_lines='skip')
            print("DataFrame carregado com sucesso:")
            print(df)
            return df
        else:
            return None
    except Exception as e:
        print(f"Erro ao carregar o arquivo CSV: {e}")
        return None

# Analisar as notas fiscais
def analisar_notas_fiscais(df):
    try:
        # Total de vendas
        total_vendas = df['valor_total'].sum()
        # Produto mais vendido
        produto_mais_vendido = df['produto'].value_counts().idxmax()
        # Total de vendas por produto
        vendas_por_produto = df.groupby('produto')['valor_total'].sum()
        # Quantidade de cada produto no estoque
        quantidade_por_produto = df.groupby('produto')['quantidade'].sum()
        return total_vendas, produto_mais_vendido, vendas_por_produto, quantidade_por_produto
    except Exception as e:
        print(f"Erro ao analisar as notas fiscais: {e}")
        return None, None, None, None

# Gerar relatório em Word
def gerar_relatorio(total_vendas, produto_mais_vendido, vendas_por_produto, quantidade_por_produto):
    try:
        print(f"Total de Vendas: R${total_vendas:.2f}")
        print(f"Produto Mais Vendido: {produto_mais_vendido}")
        print("Vendas por Produto:")
        print(vendas_por_produto)
        print("Quantidade por Produto no Estoque:")
        print(quantidade_por_produto)

        # Criar documento Word
        doc = Document()
        doc.add_heading('Relatório de Vendas', 0)
        doc.add_paragraph(f"Total de Vendas: R${total_vendas:.2f}")
        doc.add_paragraph(f"Produto Mais Vendido: {produto_mais_vendido}")
        
        doc.add_heading('Vendas por Produto:', level=1)
        for produto, valor in vendas_por_produto.items():
            doc.add_paragraph(f"{produto}: R${valor:.2f}")
        
        doc.add_heading('Quantidade por Produto no Estoque:', level=1)
        for produto, quantidade in quantidade_por_produto.items():
            doc.add_paragraph(f"{produto}: {quantidade}")
        
        # Salvar documento Word
        doc.save('relatorio de vendas.docx')
        print("Relatório gerado com sucesso.")
    except Exception as e:
        print(f"Erro ao gerar o relatório: {e}")

# Caminho do arquivo CSV
caminho_arquivo = 'D:/Programa python para FC do marcones/nota_fiscal.csv'

# Verificar o diretório de trabalho atual
print("Diretório de trabalho atual:", os.getcwd())

# Carregar e analisar as notas fiscais
df = carregar_nota_fiscal(caminho_arquivo)
if df is not None:
    total_vendas, produto_mais_vendido, vendas_por_produto, quantidade_por_produto = analisar_notas_fiscais(df)
    if total_vendas is not None:
        # Gerar o relatório em Word
        gerar_relatorio(total_vendas, produto_mais_vendido, vendas_por_produto, quantidade_por_produto)